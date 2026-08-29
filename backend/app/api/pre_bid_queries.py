from datetime import date
import csv
import io
from fastapi import APIRouter,Depends,Header,HTTPException,Query,Request
from fastapi.responses import StreamingResponse
from docx import Document as WordDocument
from docx.enum.section import WD_ORIENT
from docx.shared import Inches,Pt
from sqlalchemy import select
from sqlalchemy.orm import Session
from app.database.session import get_db
from app.models import AuditEvent,BidPreBidQuery,BidPreBidQueryDecision,BidProject,User
from app.schemas.pre_bid_queries import PreBidQueryCreate,PreBidQueryRead,PreBidQuerySuggestionDecision,PreBidQueryUpdate
from app.security.auth import Permission,current_user,require_project_access
from app.services.pre_bid_queries import approve_pre_bid_query,create_pre_bid_query,list_pre_bid_queries,pre_bid_query_summary,update_pre_bid_query
from app.services.pre_bid_query_intelligence import suggest_pre_bid_queries

router=APIRouter()
def user_dep(db:Session=Depends(get_db),x_user_id:int=Header(default=1,alias="X-User-ID")):return current_user(db,x_user_id)
def metadata(request:Request):return {"ip":request.client.host if request.client else None,"user_agent":request.headers.get("user-agent")}
def get_bid(db:Session,bid_id:int):
 bid=db.get(BidProject,bid_id)
 if not bid:raise HTTPException(404,"Bid project not found")
 return bid
def get_item(db:Session,item_id:int):
 item=db.get(BidPreBidQuery,item_id)
 if not item:raise HTTPException(404,"Pre-bid query not found")
 return item

@router.post("/bids/{bid_id}/pre-bid-queries",response_model=PreBidQueryRead,status_code=201)
def add_pre_bid_query(bid_id:int,payload:PreBidQueryCreate,request:Request,db:Session=Depends(get_db),user:User=Depends(user_dep)):
 require_project_access(db,user,bid_id,Permission.PRE_BID_QUERY_MANAGE);get_bid(db,bid_id);return create_pre_bid_query(db,bid_id,payload,user.id,metadata(request))

@router.get("/bids/{bid_id}/pre-bid-queries")
def pre_bid_queries(bid_id:int,db:Session=Depends(get_db),user:User=Depends(user_dep),search:str|None=None,query_category:str|None=None,priority:str|None=None,status:str|None=None,responsible_function:str|None=None,requirement_id:int|None=None,missing_input_id:int|None=None,source_document_id:int|None=None,target_response_from:date|None=None,target_response_to:date|None=None,page:int=Query(1,ge=1),page_size:int=Query(20,ge=1,le=100)):
 require_project_access(db,user,bid_id,Permission.PRE_BID_QUERY_VIEW);get_bid(db,bid_id);rows,total=list_pre_bid_queries(db,bid_id,locals(),page,page_size);return {"items":[PreBidQueryRead.model_validate(x).model_dump(mode="json") for x in rows],"total":total,"page":page,"page_size":page_size,"summary":pre_bid_query_summary(db,bid_id)}

@router.get("/bids/{bid_id}/pre-bid-queries/export.csv")
def export_pre_bid_queries(bid_id:int,db:Session=Depends(get_db),user:User=Depends(user_dep)):
 require_project_access(db,user,bid_id,Permission.PRE_BID_QUERY_VIEW);bid=get_bid(db,bid_id)
 rows=db.scalars(select(BidPreBidQuery).where(BidPreBidQuery.bid_project_id==bid_id,BidPreBidQuery.status!="Withdrawn").order_by(BidPreBidQuery.id)).all()
 output=io.StringIO(newline="")
 writer=csv.writer(output)
 writer.writerow(["Query No.","Query Title","Tender Reference / Source","Clause","Page","Pre-Bid Query","Category","Priority","Status","Responsible Function","Target Response Date","Employer Response","Response Reference","Response Date"])
 for item in rows:
  writer.writerow([
   item.query_number or item.id,item.query_title,item.source_document_title or item.source_original_filename or "",
   item.source_clause or "",item.source_page or "",item.query_text,item.query_category,item.priority,item.status,
   item.responsible_function or "",item.target_response_date.isoformat() if item.target_response_date else "",
   item.employer_response or "",item.response_reference or "",item.response_date.isoformat() if item.response_date else "",
  ])
 body="\ufeff"+output.getvalue()
 filename=f"{bid.bid_id}_pre_bid_queries.csv"
 return StreamingResponse(iter([body]),media_type="text/csv; charset=utf-8",headers={"Content-Disposition":f'attachment; filename="{filename}"'})

@router.get("/bids/{bid_id}/pre-bid-queries/export.docx")
def export_pre_bid_queries_docx(bid_id:int,submission_only:bool=Query(False),db:Session=Depends(get_db),user:User=Depends(user_dep)):
 require_project_access(db,user,bid_id,Permission.PRE_BID_QUERY_VIEW);bid=get_bid(db,bid_id)
 conditions=[BidPreBidQuery.bid_project_id==bid_id,BidPreBidQuery.status!="Withdrawn"]
 if submission_only:conditions.append(BidPreBidQuery.status.in_(["Ready for Review","Submitted","Responded","Closed"]))
 rows=db.scalars(select(BidPreBidQuery).where(*conditions).order_by(BidPreBidQuery.query_number.asc().nullslast(),BidPreBidQuery.id)).all()
 document=WordDocument()
 section=document.sections[0]
 section.orientation=WD_ORIENT.LANDSCAPE
 section.page_width,section.page_height=section.page_height,section.page_width
 section.top_margin=Inches(.55);section.bottom_margin=Inches(.55);section.left_margin=Inches(.55);section.right_margin=Inches(.55)
 styles=document.styles
 styles["Normal"].font.name="Calibri";styles["Normal"].font.size=Pt(9)
 title=document.add_paragraph()
 title_run=title.add_run("PRE-BID QUERY REGISTER" + (" — SUBMISSION SET" if submission_only else ""));title_run.bold=True;title_run.font.size=Pt(16)
 subtitle=document.add_paragraph()
 subtitle.add_run(f"Bid ID: {bid.bid_id}   |   Tender Ref: {bid.tender_reference_no}\n").bold=True
 subtitle.add_run(f"Tender: {bid.tender_name}\nClient: {bid.client}")
 table=document.add_table(rows=1,cols=6)
 table.style="Table Grid"
 headers=["Query No.","Tender Reference","Clause / Page","Pre-Bid Query","Category / Priority","Employer Response"]
 for i,label in enumerate(headers):
  cell=table.rows[0].cells[i];cell.text=label
  for run in cell.paragraphs[0].runs:run.bold=True
 for item in rows:
  cells=table.add_row().cells
  cells[0].text=str(item.query_number or item.id)
  cells[1].text=item.source_document_title or item.source_original_filename or ""
  location=[]
  if item.source_clause:location.append(f"Cl. {item.source_clause}")
  if item.source_page:location.append(f"p. {item.source_page}")
  cells[2].text=" / ".join(location)
  cells[3].text=item.query_text
  cells[4].text=f"{item.query_category}\n{item.priority}\n{item.status}"
  response=item.employer_response or ""
  if item.response_reference:response+=f"\nRef: {item.response_reference}"
  if item.response_date:response+=f"\nDate: {item.response_date.isoformat()}"
  cells[5].text=response
 document.add_paragraph(f"Generated from L&T Bid Intelligence · {len(rows)} query{'ies' if len(rows)!=1 else ''}" + (" · Draft and withdrawn items excluded" if submission_only else ""))
 output=io.BytesIO();document.save(output);output.seek(0)
 filename=f"{bid.bid_id}_pre_bid_queries{'_submission' if submission_only else ''}.docx"
 return StreamingResponse(output,media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",headers={"Content-Disposition":f'attachment; filename="{filename}"'})

@router.post("/bids/{bid_id}/pre-bid-query-suggestions/decision")
def decide_pre_bid_query_suggestion(bid_id:int,payload:PreBidQuerySuggestionDecision,request:Request,db:Session=Depends(get_db),user:User=Depends(user_dep)):
 require_project_access(db,user,bid_id,Permission.PRE_BID_QUERY_MANAGE);get_bid(db,bid_id)
 item=db.scalar(select(BidPreBidQueryDecision).where(
  BidPreBidQueryDecision.bid_project_id==bid_id,
  BidPreBidQueryDecision.source_kind==payload.source_kind,
  BidPreBidQueryDecision.source_id==payload.source_id,
 ))
 if payload.decision=="Reconsider":
  if item:db.delete(item)
  event="pre_bid_query_suggestion.reconsidered"
 else:
  if item:
   item.decision=payload.decision;item.reason=payload.reason;item.decided_by=user.id
  else:
   item=BidPreBidQueryDecision(
    bid_project_id=bid_id,source_kind=payload.source_kind,source_id=payload.source_id,
    decision=payload.decision,reason=payload.reason,decided_by=user.id,
   )
   db.add(item)
  event="pre_bid_query_suggestion.suppressed"
 db.add(AuditEvent(
  user_id=user.id,bid_project_id=bid_id,event_type=event,entity_type="PreBidQuerySuggestion",
  entity_id=f"{payload.source_kind}:{payload.source_id}",request_metadata=metadata(request),
  details={"source_kind":payload.source_kind,"source_id":payload.source_id,"decision":payload.decision,"reason":payload.reason},
 ))
 db.commit()
 return {"source_kind":payload.source_kind,"source_id":payload.source_id,"decision":payload.decision}

@router.get("/bids/{bid_id}/pre-bid-query-suggestions")
def pre_bid_query_suggestions(bid_id:int,db:Session=Depends(get_db),user:User=Depends(user_dep)):
 require_project_access(db,user,bid_id,Permission.PRE_BID_QUERY_VIEW);get_bid(db,bid_id);return suggest_pre_bid_queries(db,bid_id)

@router.post("/pre-bid-queries/{query_id}/approve",response_model=PreBidQueryRead)
def approve_query(query_id:int,request:Request,db:Session=Depends(get_db),user:User=Depends(user_dep)):
 item=get_item(db,query_id);require_project_access(db,user,item.bid_project_id,Permission.PRE_BID_QUERY_APPROVE)
 return approve_pre_bid_query(db,item,user.id,metadata(request))

@router.get("/pre-bid-queries/{query_id}",response_model=PreBidQueryRead)
def pre_bid_query_detail(query_id:int,db:Session=Depends(get_db),user:User=Depends(user_dep)):
 item=get_item(db,query_id);require_project_access(db,user,item.bid_project_id,Permission.PRE_BID_QUERY_VIEW);return item

@router.patch("/pre-bid-queries/{query_id}",response_model=PreBidQueryRead)
def edit_pre_bid_query(query_id:int,payload:PreBidQueryUpdate,request:Request,db:Session=Depends(get_db),user:User=Depends(user_dep)):
 item=get_item(db,query_id);require_project_access(db,user,item.bid_project_id,Permission.PRE_BID_QUERY_MANAGE);return update_pre_bid_query(db,item,payload,user.id,metadata(request))
